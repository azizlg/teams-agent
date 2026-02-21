"""
exporters.py — DOCX & PDF export for meeting reports.

Converts generated report content into downloadable DOCX and PDF files.
Uses python-docx for Word documents and WeasyPrint for PDF generation.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any

from reports.templates import ReportTemplate, get_template

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Base exporter
# ---------------------------------------------------------------------------

class BaseExporter:
    """Base class for report exporters."""

    def __init__(self, template: ReportTemplate | None = None) -> None:
        self._template = template or get_template("full_report")

    async def export(
        self,
        content: str,
        output_path: str | Path | None = None,
    ) -> bytes:
        """
        Export content to the target format.

        Args:
            content: Rendered report content (HTML or Markdown).
            output_path: Optional path to write the file.

        Returns:
            File bytes.
        """
        raise NotImplementedError


# ---------------------------------------------------------------------------
# DOCX exporter
# ---------------------------------------------------------------------------

class DocxExporter(BaseExporter):
    """
    Export reports as Microsoft Word (.docx) documents.

    Uses python-docx to build a styled Word document from
    the report's structured section data.
    """

    async def export(
        self,
        content: str,
        output_path: str | Path | None = None,
        *,
        sections: list[dict[str, Any]] | None = None,
        meeting_title: str = "",
        meeting_date: str = "",
    ) -> bytes:
        """
        Export to DOCX format.

        Args:
            content: Markdown content (used as fallback).
            output_path: Optional file path.
            sections: Structured section data for precise rendering.
            meeting_title: Meeting title for the header.
            meeting_date: Meeting date for the subtitle.

        Returns:
            DOCX file bytes.
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # ---- Document styles ----
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # ---- Title ----
        title_para = doc.add_heading(self._template.header_title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---- Subtitle ----
        subtitle_text = self._template.subtitle_template.format(
            meeting_title=meeting_title or "Untitled Meeting",
            meeting_date=meeting_date or "Date not specified",
        )
        subtitle = doc.add_paragraph(subtitle_text)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")  # spacer

        # ---- Sections ----
        if sections:
            for sec in sections:
                if not sec.get("has_data"):
                    continue

                # Section heading
                heading_text = f"{sec.get('icon', '')} {sec['title']}".strip()
                doc.add_heading(heading_text, level=1)

                # Section content
                data = sec.get("data")
                self._render_data_to_docx(doc, data)

                doc.add_paragraph("")  # spacer
        else:
            # Fallback: render raw markdown content
            for line in content.split("\n"):
                line = line.strip()
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("*"))
                    run.bold = True
                elif line:
                    doc.add_paragraph(line)

        # ---- Footer ----
        doc.add_paragraph("")
        footer = doc.add_paragraph(self._template.footer_text)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.style.font.size = Pt(9)
        footer.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ---- Save ----
        buffer = io.BytesIO()
        doc.save(buffer)
        file_bytes = buffer.getvalue()

        if output_path:
            Path(output_path).write_bytes(file_bytes)
            logger.info("DOCX exported to %s (%d bytes)", output_path, len(file_bytes))

        return file_bytes

    def _render_data_to_docx(self, doc: Any, data: Any) -> None:
        """Render structured data into a DOCX document."""
        from docx.shared import Pt

        if isinstance(data, str):
            doc.add_paragraph(data)

        elif isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    first_key = next(iter(item), None)
                    if first_key:
                        # First field as bold
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(f"{first_key.replace('_', ' ').title()}: ")
                        run.bold = True
                        p.add_run(str(item[first_key]))

                        # Remaining fields as sub-items
                        for k, v in item.items():
                            if k == first_key or k.startswith("_"):
                                continue
                            display_k = k.replace("_", " ").title()
                            if isinstance(v, list):
                                v = ", ".join(str(x) for x in v)
                            sub_p = doc.add_paragraph(
                                f"{display_k}: {v}",
                                style="List Bullet 2",
                            )
                else:
                    doc.add_paragraph(str(item), style="List Bullet")

        elif isinstance(data, dict):
            for k, v in data.items():
                if k.startswith("_"):
                    continue
                display_k = k.replace("_", " ").title()
                if isinstance(v, str):
                    p = doc.add_paragraph()
                    run = p.add_run(f"{display_k}: ")
                    run.bold = True
                    p.add_run(v)
                elif isinstance(v, list):
                    doc.add_heading(display_k, level=2)
                    self._render_data_to_docx(doc, v)
                elif isinstance(v, dict):
                    doc.add_heading(display_k, level=2)
                    self._render_data_to_docx(doc, v)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{display_k}: ")
                    run.bold = True
                    p.add_run(str(v))


# ---------------------------------------------------------------------------
# PDF exporter
# ---------------------------------------------------------------------------

class PdfExporter(BaseExporter):
    """
    Export reports as PDF documents.

    Uses WeasyPrint to convert HTML content to PDF.
    """

    async def export(
        self,
        content: str,
        output_path: str | Path | None = None,
        *,
        html_content: str | None = None,
    ) -> bytes:
        """
        Export to PDF format.

        Args:
            content: HTML content (or Markdown to convert).
            output_path: Optional file path.
            html_content: Pre-rendered HTML (preferred over content).

        Returns:
            PDF file bytes.
        """
        import asyncio
        from weasyprint import HTML

        html = html_content or content

        # If content looks like Markdown, wrap in basic HTML
        if not html.strip().startswith("<!DOCTYPE") and not html.strip().startswith("<html"):
            html = self._markdown_to_html(html)

        # WeasyPrint is CPU-bound, run in executor
        loop = asyncio.get_running_loop()
        pdf_bytes = await loop.run_in_executor(
            None,
            lambda: HTML(string=html).write_pdf(),
        )

        if output_path:
            Path(output_path).write_bytes(pdf_bytes)
            logger.info("PDF exported to %s (%d bytes)", output_path, len(pdf_bytes))

        return pdf_bytes

    def _markdown_to_html(self, md_content: str) -> str:
        """Convert Markdown content to styled HTML for PDF rendering."""
        lines = md_content.split("\n")
        html_lines: list[str] = [
            "<!DOCTYPE html>",
            "<html><head>",
            f"<style>",
            f"  body {{ font-family: {self._template.font_family}; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 2rem; }}",
            f"  h1 {{ color: {self._template.primary_color}; border-bottom: 2px solid {self._template.primary_color}; }}",
            f"  h2 {{ color: {self._template.accent_color}; }}",
            f"</style>",
            "</head><body>",
        ]

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("# "):
                html_lines.append(f"<h1>{stripped[2:]}</h1>")
            elif stripped.startswith("## "):
                html_lines.append(f"<h2>{stripped[3:]}</h2>")
            elif stripped.startswith("### "):
                html_lines.append(f"<h3>{stripped[4:]}</h3>")
            elif stripped.startswith("- "):
                html_lines.append(f"<li>{stripped[2:]}</li>")
            elif stripped.startswith("**") and stripped.endswith("**"):
                html_lines.append(f"<p><strong>{stripped.strip('*')}</strong></p>")
            elif stripped:
                html_lines.append(f"<p>{stripped}</p>")

        html_lines.extend(["</body>", "</html>"])
        return "\n".join(html_lines)


# ---------------------------------------------------------------------------
# Export dispatcher
# ---------------------------------------------------------------------------

async def export_report(
    content: str,
    output_format: str,
    output_path: str | Path | None = None,
    *,
    sections: list[dict[str, Any]] | None = None,
    html_content: str | None = None,
    meeting_title: str = "",
    meeting_date: str = "",
    template_name: str = "full_report",
) -> bytes:
    """
    Convenience function to export a report in the specified format.

    Args:
        content: Report content (Markdown or HTML).
        output_format: 'docx' or 'pdf'.
        output_path: Optional output file path.
        sections: Structured section data.
        html_content: Pre-rendered HTML for PDF.
        meeting_title: Meeting title.
        meeting_date: Meeting date.
        template_name: Template name.

    Returns:
        File bytes.
    """
    template = get_template(template_name)

    if output_format == "docx":
        exporter = DocxExporter(template=template)
        return await exporter.export(
            content,
            output_path,
            sections=sections,
            meeting_title=meeting_title,
            meeting_date=meeting_date,
        )
    elif output_format == "pdf":
        exporter = PdfExporter(template=template)
        return await exporter.export(
            content,
            output_path,
            html_content=html_content,
        )
    else:
        raise ValueError(f"Unsupported format: {output_format}. Use 'docx' or 'pdf'.")


# ---------------------------------------------------------------------------
# Unit test stub
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import asyncio

    async def _test() -> None:
        # Test DOCX exporter can be instantiated
        docx_exporter = DocxExporter()
        assert docx_exporter._template.name == "full_report"

        # Test PDF exporter can be instantiated
        pdf_exporter = PdfExporter()
        assert pdf_exporter._template.name == "full_report"

        # Test markdown to HTML conversion
        md = "# Title\n\n## Section\n\n- Item 1\n- Item 2"
        html = pdf_exporter._markdown_to_html(md)
        assert "<h1>Title</h1>" in html
        assert "<h2>Section</h2>" in html
        assert "<li>Item 1</li>" in html

        print("✓ Exporter tests passed")
        print("  (Full export tests require python-docx and weasyprint)")

    asyncio.run(_test())
